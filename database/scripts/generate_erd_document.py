from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATABASE_DIR = ROOT / "database"
OUTPUT_DOCX = DATABASE_DIR / "Rancangan_ERD_VAULTLAUNDRY.docx"
DIAGRAM_PNG = DATABASE_DIR / "erd-vaultlaundry.png"
MERMAID_FILE = DATABASE_DIR / "erd-vaultlaundry.mmd"

ORANGE = "FF6626"
ORANGE_LIGHT = "FFF0E8"
INK = "181512"
MUTED = "6B625A"
BORDER = "E4D8CA"
CREAM = "FFF9F1"
WHITE = "FFFFFF"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"

FONT_NAME = "Calibri"
MONO_FONT = "Consolas"

ENTITY_DATA = {
    "users": [
        ("id", "BIGINT", "PK", "Tidak", "Primary key."),
        ("name", "VARCHAR(255)", "-", "Tidak", "Nama pengguna."),
        ("email", "VARCHAR(255)", "UK", "Tidak", "Email login, harus unik."),
        ("email_verified_at", "TIMESTAMP", "-", "Ya", "Waktu verifikasi email."),
        ("password", "VARCHAR(255)", "-", "Tidak", "Password yang telah di-hash."),
        ("role", "VARCHAR(255)", "-", "Tidak", "Default user; nilai aplikasi: admin, kasir, user."),
        ("remember_token", "VARCHAR(100)", "-", "Ya", "Token remember-me Laravel."),
        ("created_at", "TIMESTAMP", "-", "Ya", "Waktu pembuatan record."),
        ("updated_at", "TIMESTAMP", "-", "Ya", "Waktu pembaruan record."),
    ],
    "customers": [
        ("id", "BIGINT", "PK", "Tidak", "Primary key."),
        ("user_id", "BIGINT", "FK", "Ya", "Referensi users.id; nullOnDelete."),
        ("name", "VARCHAR(255)", "-", "Tidak", "Nama pelanggan."),
        ("phone", "VARCHAR(255)", "-", "Ya", "Nomor telepon pelanggan."),
        ("address", "TEXT", "-", "Ya", "Alamat pelanggan."),
        ("gender", "ENUM", "-", "Ya", "male atau female."),
        ("notes", "TEXT", "-", "Ya", "Catatan khusus pelanggan."),
        ("created_at", "TIMESTAMP", "-", "Ya", "Waktu pembuatan record."),
        ("updated_at", "TIMESTAMP", "-", "Ya", "Waktu pembaruan record."),
    ],
    "services": [
        ("id", "BIGINT", "PK", "Tidak", "Primary key."),
        ("name", "VARCHAR(255)", "-", "Tidak", "Nama layanan laundry."),
        ("description", "TEXT", "-", "Ya", "Deskripsi layanan."),
        ("price_per_kg", "DECIMAL(10,2)", "-", "Tidak", "Tarif layanan per kilogram."),
        ("estimated_days", "INTEGER", "-", "Tidak", "Estimasi hari, default 2."),
        ("is_active", "BOOLEAN", "-", "Tidak", "Status layanan, default true."),
        ("created_at", "TIMESTAMP", "-", "Ya", "Waktu pembuatan record."),
        ("updated_at", "TIMESTAMP", "-", "Ya", "Waktu pembaruan record."),
    ],
    "bookings": [
        ("id", "BIGINT", "PK", "Tidak", "Primary key."),
        ("booking_code", "VARCHAR(255)", "UK", "Tidak", "Kode booking unik LDY-YYYY-0001."),
        ("user_id", "BIGINT", "FK", "Ya", "Referensi users.id; nullOnDelete."),
        ("customer_id", "BIGINT", "FK", "Ya", "Referensi customers.id; nullOnDelete."),
        ("service_id", "BIGINT", "FK", "Tidak", "Referensi services.id; cascadeOnDelete."),
        ("booking_date", "DATE", "-", "Tidak", "Tanggal booking."),
        ("estimated_finish_date", "DATE", "-", "Ya", "Tanggal estimasi selesai."),
        ("weight", "DECIMAL(8,2)", "-", "Ya", "Berat laundry dalam kilogram."),
        ("total_price", "DECIMAL(12,2)", "-", "Tidak", "Total harga, default 0."),
        ("pickup_type", "ENUM", "-", "Tidak", "antar_sendiri atau pickup."),
        ("status", "ENUM", "-", "Tidak", "Tahapan proses laundry."),
        ("notes", "TEXT", "-", "Ya", "Catatan booking."),
        ("created_at", "TIMESTAMP", "-", "Ya", "Waktu pembuatan record."),
        ("updated_at", "TIMESTAMP", "-", "Ya", "Waktu pembaruan record."),
    ],
    "payments": [
        ("id", "BIGINT", "PK", "Tidak", "Primary key."),
        ("booking_id", "BIGINT", "FK", "Tidak", "Referensi bookings.id; cascadeOnDelete."),
        ("payment_code", "VARCHAR(255)", "UK", "Tidak", "Kode pembayaran unik PAY-YYYY-0001."),
        ("payment_date", "DATETIME", "-", "Tidak", "Waktu transaksi pembayaran."),
        ("payment_method", "ENUM", "-", "Tidak", "cash, transfer, atau ewallet."),
        ("amount_paid", "DECIMAL(12,2)", "-", "Tidak", "Jumlah yang dibayar."),
        ("total_bill", "DECIMAL(12,2)", "-", "Tidak", "Total tagihan dari booking."),
        ("change_amount", "DECIMAL(12,2)", "-", "Tidak", "Kembalian, default 0."),
        ("payment_status", "ENUM", "-", "Tidak", "unpaid, partial, atau paid."),
        ("notes", "TEXT", "-", "Ya", "Catatan pembayaran."),
        ("processed_by", "BIGINT", "FK", "Ya", "Referensi users.id; nullOnDelete."),
        ("created_at", "TIMESTAMP", "-", "Ya", "Waktu pembuatan record."),
        ("updated_at", "TIMESTAMP", "-", "Ya", "Waktu pembaruan record."),
    ],
}

ENTITY_DESCRIPTIONS = {
    "users": "Menyimpan akun autentikasi dan role admin, kasir, atau user.",
    "customers": "Menyimpan profil pelanggan, termasuk pelanggan walk-in tanpa akun.",
    "services": "Menyimpan master layanan, harga per kilogram, dan estimasi hari.",
    "bookings": "Menyimpan pesanan laundry, layanan yang dipilih, nilai transaksi, dan status proses.",
    "payments": "Menyimpan transaksi pembayaran dan pengguna yang memproses pembayaran.",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_fixed_table_layout(table, widths_inches: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_dxa = int(sum(widths_inches) * 1440)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_dxa = int(widths_inches[index] * 1440)
            cell.width = Inches(widths_inches[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_paragraph_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def set_repeatable_font(run, name=FONT_NAME, size=10, bold=False, color=INK, italic=False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    set_repeatable_font(run, size=8.5, color=MUTED)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def apply_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, ORANGE, 14, 7),
        ("Heading 2", 13, INK, 11, 5),
        ("Heading 3", 11, MUTED, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("ERD Code", 1)
    code.font.name = MONO_FONT
    code._element.rPr.rFonts.set(qn("w:ascii"), MONO_FONT)
    code._element.rPr.rFonts.set(qn("w:hAnsi"), MONO_FONT)
    code.font.size = Pt(7.3)
    code.font.color.rgb = rgb(INK)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1


def setup_section(section, landscape=False) -> None:
    section.different_first_page_header_footer = False
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)


def add_running_furniture(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("VAULTLAUNDRY  /  RANCANGAN ENTITY RELATIONSHIP DIAGRAM")
    set_repeatable_font(run, size=8, bold=True, color=MUTED)
    p.paragraph_format.space_after = Pt(2)

    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), ORANGE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    add_page_number(footer.paragraphs[0])


def add_title_block(doc: Document) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(10)
    run = kicker.add_run("DOKUMENTASI PERANCANGAN DATABASE")
    set_repeatable_font(run, size=10, bold=True, color=ORANGE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Entity Relationship Diagram (ERD)")
    set_repeatable_font(run, size=26, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run("Sistem Informasi Laundry VAULTLAUNDRY")
    set_repeatable_font(run, size=18, bold=True, color=ORANGE)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_after = Pt(28)
    run = desc.add_run("Aplikasi berbasis Laravel dengan PostgreSQL sebagai database utama")
    set_repeatable_font(run, size=11, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("NAMA", "[Isi Nama]"),
        ("NIM", "[Isi NIM]"),
        ("KELAS", "[Isi Kelas]"),
        ("MATA KULIAH", "Pemrograman Web Lanjut"),
    ]
    for row, values in zip(table.rows, rows):
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = ""
            set_cell_shading(cell, ORANGE_LIGHT if index == 0 else CREAM)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_repeatable_font(run, size=10.5, bold=(index == 0), color=INK if index else ORANGE)
    set_fixed_table_layout(table, [1.65, 4.45])
    set_table_borders(table)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Disusun berdasarkan migration dan model aktif pada project VAULTLAUNDRY.")
    set_repeatable_font(run, size=9, italic=True, color=MUTED)

    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, fill=ORANGE_LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}: ")
    set_repeatable_font(label_run, size=10, bold=True, color=ORANGE)
    text_run = p.add_run(text)
    set_repeatable_font(text_run, size=10, color=INK)
    set_fixed_table_layout(table, [6.65])
    set_table_borders(table, color="F0C9B8")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_overview(doc: Document) -> None:
    doc.add_heading("A. Pendahuluan", level=1)
    p = doc.add_paragraph(
        "VAULTLAUNDRY merupakan aplikasi pengelolaan operasional laundry berbasis Laravel. "
        "Sistem mendukung autentikasi berbasis role, pengelolaan pelanggan dan layanan, "
        "booking laundry, monitoring status, pembayaran, invoice PDF, dashboard realtime, "
        "serta laporan transaksi dan pendapatan."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph(
        "Dokumen ini memodelkan lima entitas bisnis inti yang benar-benar memiliki tabel fisik, "
        "yaitu users, customers, services, bookings, dan payments. Tabel pendukung Laravel "
        "seperti sessions, cache, jobs, job_batches, dan failed_jobs tidak ditampilkan pada ERD "
        "inti karena tidak merepresentasikan proses bisnis laundry."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_callout(
        doc,
        "Sumber kebenaran",
        "Struktur atribut, foreign key, nullable field, unique constraint, dan aksi ON DELETE "
        "mengikuti file Laravel migration pada database/migrations. schema.sql hanya dokumentasi/manual import alternatif.",
    )

    doc.add_heading("B. Ruang Lingkup ERD", level=1)
    items = [
        "Akun dan role pengguna melalui tabel users.",
        "Profil pelanggan melalui tabel customers.",
        "Master layanan laundry melalui tabel services.",
        "Pesanan dan status proses laundry melalui tabel bookings.",
        "Pembayaran dan petugas pemroses melalui tabel payments.",
        "Invoice dan report tidak mempunyai tabel fisik; keduanya dihasilkan dari query bookings dan payments.",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.add_heading("C. Notasi yang Digunakan", level=1)
    table = doc.add_table(rows=1, cols=2)
    headers = ("Notasi", "Arti")
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
        set_cell_shading(table.cell(0, i), ORANGE)
    notation_rows = [
        ("PK", "Primary Key, identitas unik setiap record."),
        ("FK", "Foreign Key, penghubung ke primary key tabel lain."),
        ("UK", "Unique Key, nilai wajib unik."),
        ("0..1", "Relasi bersifat opsional dan maksimal satu."),
        ("0..N", "Relasi bersifat opsional dan dapat berjumlah banyak."),
        ("1", "Relasi wajib tepat satu."),
    ]
    for notation, meaning in notation_rows:
        cells = table.add_row().cells
        cells[0].text = notation
        cells[1].text = meaning
    set_fixed_table_layout(table, [1.2, 5.45])
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    style_table_text(table)


def style_table_text(table) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_repeatable_font(
                        run,
                        size=8.7 if row_index else 9,
                        bold=(row_index == 0),
                        color=WHITE if row_index == 0 else INK,
                    )


def add_entity_sections(doc: Document) -> None:
    doc.add_heading("1. Daftar Entitas dan Atribut", level=1)
    p = doc.add_paragraph(
        "Daftar berikut menyajikan atribut berdasarkan migration aktif. Kolom nullable/default "
        "ditulis sesuai batasan database, bukan hanya berdasarkan form aplikasi."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for index, (entity, rows) in enumerate(ENTITY_DATA.items(), 1):
        heading = doc.add_heading(f"1.{index} Tabel {entity}", level=2)
        keep_paragraph_with_next(heading)
        p = doc.add_paragraph(ENTITY_DESCRIPTIONS[entity])
        p.paragraph_format.space_after = Pt(5)
        keep_paragraph_with_next(p)

        table = doc.add_table(rows=1, cols=5)
        headers = ("Field", "Tipe Data", "Key", "Nullable", "Keterangan")
        for column, header in enumerate(headers):
            cell = table.cell(0, column)
            cell.text = header
            set_cell_shading(cell, ORANGE)
        for row_values in rows:
            cells = table.add_row().cells
            for column, value in enumerate(row_values):
                cells[column].text = value
                if len(table.rows) % 2 == 0:
                    set_cell_shading(cells[column], CREAM)
        set_fixed_table_layout(table, [1.15, 1.35, 0.48, 0.68, 3.0])
        set_table_borders(table)
        repeat_table_header(table.rows[0])
        style_table_text(table)

        if index < len(ENTITY_DATA):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(1)


def add_relationships(doc: Document) -> None:
    doc.add_heading("2. Hubungan Antar Entitas", level=1)
    relationships = [
        (
            "Users ke Customers",
            "0..1 : 0..1 secara logis",
            "Satu akun user memiliki maksimal satu profil customer melalui relasi Eloquent hasOne. "
            "Kolom customers.user_id nullable sehingga customer walk-in boleh tidak mempunyai akun.",
            "customers.user_id -> users.id; ON DELETE SET NULL.",
        ),
        (
            "Users ke Bookings",
            "0..1 : 0..N",
            "Satu user dapat membuat banyak booking. Booking kasir untuk pelanggan tanpa akun dapat memiliki user_id null.",
            "bookings.user_id -> users.id; ON DELETE SET NULL.",
        ),
        (
            "Customers ke Bookings",
            "0..1 : 0..N",
            "Satu customer dapat mempunyai banyak booking. Foreign key nullable menjaga histori booking ketika profil customer dihapus.",
            "bookings.customer_id -> customers.id; ON DELETE SET NULL.",
        ),
        (
            "Services ke Bookings",
            "1 : 0..N",
            "Setiap booking wajib memilih tepat satu service. Satu service dapat digunakan oleh banyak booking.",
            "bookings.service_id -> services.id; ON DELETE CASCADE.",
        ),
        (
            "Bookings ke Payments",
            "1 : 0..1 secara logis",
            "Model Booking menggunakan hasOne Payment dan alur aplikasi memperlakukan satu booking sebagai satu pembayaran.",
            "payments.booking_id -> bookings.id; ON DELETE CASCADE.",
        ),
        (
            "Users ke Payments",
            "0..1 : 0..N",
            "Admin atau kasir dapat memproses banyak pembayaran. processed_by boleh null untuk pembayaran yang belum diproses.",
            "payments.processed_by -> users.id; ON DELETE SET NULL.",
        ),
    ]

    for index, (name, cardinality, explanation, fk) in enumerate(relationships, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        name_run = p.add_run(f"{index}. {name} ({cardinality})")
        set_repeatable_font(name_run, size=10.5, bold=True, color=INK)
        p = doc.add_paragraph(explanation)
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.space_after = Pt(7)
        label = p.add_run("Foreign key: ")
        set_repeatable_font(label, size=9.5, bold=True, color=ORANGE)
        value = p.add_run(fk)
        set_repeatable_font(value, size=9.5, color=MUTED)

    add_callout(
        doc,
        "Catatan constraint fisik",
        "customers.user_id dan payments.booking_id belum memiliki UNIQUE constraint pada migration. "
        "Karena itu database secara fisik masih dapat menerima lebih dari satu customer per user atau lebih dari satu payment per booking, "
        "walaupun model Eloquent mendefinisikan hasOne. Jika aturan 1:1 harus dijamin database, tambahkan unique index melalui migration baru.",
        fill=LIGHT_BLUE,
    )


def get_font(size: int, bold=False):
    names = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for name in names:
        if name.exists():
            return ImageFont.truetype(str(name), size=size)
    return ImageFont.load_default()


def rounded_box(draw, box, radius, fill, outline, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_entity(draw, x, y, width, title, rows, accent):
    title_height = 66
    row_height = 39
    height = title_height + row_height * len(rows) + 18
    rounded_box(draw, (x, y, x + width, y + height), 18, "#FFFDF9", accent, 4)
    draw.rounded_rectangle((x, y, x + width, y + title_height), radius=18, fill=accent)
    draw.rectangle((x, y + title_height - 18, x + width, y + title_height), fill=accent)
    title_font = get_font(27, bold=True)
    row_font = get_font(21)
    key_font = get_font(18, bold=True)
    draw.text((x + 20, y + 17), title.upper(), fill="white", font=title_font)

    current_y = y + title_height
    for index, (field, data_type, key) in enumerate(rows):
        if index % 2:
            draw.rectangle((x + 2, current_y, x + width - 2, current_y + row_height), fill="#FFF7F0")
        draw.line((x, current_y, x + width, current_y), fill="#E8DCCB", width=1)
        draw.text((x + 16, current_y + 8), field, fill="#181512", font=row_font)
        type_width = draw.textlength(data_type, font=row_font)
        if key:
            badge_width = draw.textlength(key, font=key_font) + 18
            badge_x = x + width - badge_width - 12
            type_x = badge_x - type_width - 16
            draw.text((type_x, current_y + 8), data_type, fill="#6B625A", font=row_font)
            draw.rounded_rectangle(
                (badge_x, current_y + 7, x + width - 9, current_y + 32),
                radius=8,
                fill="#FFE1D2",
            )
            draw.text((badge_x + 9, current_y + 9), key, fill="#C7440E", font=key_font)
        else:
            draw.text((x + width - type_width - 18, current_y + 8), data_type, fill="#6B625A", font=row_font)
        current_y += row_height
    return (x, y, x + width, y + height)


def draw_relation(draw, start, end, label, start_card, end_card, bend=None):
    line_color = "#72675C"
    points = [start]
    if bend:
        points.extend(bend)
    points.append(end)
    draw.line(points, fill=line_color, width=4, joint="curve")
    radius = 7
    for point in (start, end):
        draw.ellipse(
            (point[0] - radius, point[1] - radius, point[0] + radius, point[1] + radius),
            fill="#FF6626",
        )
    font = get_font(19, bold=True)
    small = get_font(18, bold=True)
    middle = points[len(points) // 2]
    label_width = draw.textlength(label, font=font)
    box = (
        middle[0] - label_width / 2 - 12,
        middle[1] - 18,
        middle[0] + label_width / 2 + 12,
        middle[1] + 18,
    )
    draw.rounded_rectangle(box, radius=8, fill="#FFF9F1", outline="#E8DCCB", width=2)
    draw.text((box[0] + 12, box[1] + 6), label, fill="#181512", font=font)

    next_point = points[1]
    start_width = draw.textlength(start_card, font=small)
    if next_point[0] < start[0]:
        start_x = start[0] - start_width - 12
    else:
        start_x = start[0] + 12
    draw.text((start_x, start[1] - 31), start_card, fill="#FF6626", font=small)

    previous_point = points[-2]
    end_width = draw.textlength(end_card, font=small)
    if end[0] < previous_point[0]:
        end_x = end[0] + 12
    else:
        end_x = end[0] - end_width - 12
    draw.text((end_x, end[1] - 31), end_card, fill="#FF6626", font=small)


def generate_diagram() -> None:
    DIAGRAM_PNG.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (2700, 1700), "#FAF4EA")
    draw = ImageDraw.Draw(image)

    title_font = get_font(43, bold=True)
    subtitle_font = get_font(23)
    draw.text((80, 42), "VAULTLAUNDRY - CORE ENTITY RELATIONSHIP DIAGRAM", fill="#181512", font=title_font)
    draw.text(
        (82, 100),
        "Kardinalitas logis aplikasi; detail constraint fisik dijelaskan pada dokumen.",
        fill="#6B625A",
        font=subtitle_font,
    )

    users = [
        ("id", "bigint", "PK"),
        ("name", "varchar", ""),
        ("email", "varchar", "UK"),
        ("email_verified_at", "timestamp?", ""),
        ("password", "varchar", ""),
        ("role", "varchar", ""),
        ("remember_token", "varchar?", ""),
        ("created_at", "timestamp?", ""),
        ("updated_at", "timestamp?", ""),
    ]
    customers = [
        ("id", "bigint", "PK"),
        ("user_id", "bigint?", "FK"),
        ("name", "varchar", ""),
        ("phone", "varchar?", ""),
        ("address", "text?", ""),
        ("gender", "enum?", ""),
        ("notes", "text?", ""),
        ("created_at", "timestamp?", ""),
        ("updated_at", "timestamp?", ""),
    ]
    services = [
        ("id", "bigint", "PK"),
        ("name", "varchar", ""),
        ("description", "text?", ""),
        ("price_per_kg", "decimal", ""),
        ("estimated_days", "integer", ""),
        ("is_active", "boolean", ""),
        ("created_at", "timestamp?", ""),
        ("updated_at", "timestamp?", ""),
    ]
    bookings = [
        ("id", "bigint", "PK"),
        ("booking_code", "varchar", "UK"),
        ("user_id", "bigint?", "FK"),
        ("customer_id", "bigint?", "FK"),
        ("service_id", "bigint", "FK"),
        ("booking_date", "date", ""),
        ("estimated_finish_date", "date?", ""),
        ("weight", "decimal?", ""),
        ("total_price", "decimal", ""),
        ("pickup_type", "enum", ""),
        ("status", "enum", ""),
        ("notes", "text?", ""),
        ("created_at", "timestamp?", ""),
        ("updated_at", "timestamp?", ""),
    ]
    payments = [
        ("id", "bigint", "PK"),
        ("booking_id", "bigint", "FK"),
        ("payment_code", "varchar", "UK"),
        ("payment_date", "datetime", ""),
        ("payment_method", "enum", ""),
        ("amount_paid", "decimal", ""),
        ("total_bill", "decimal", ""),
        ("change_amount", "decimal", ""),
        ("payment_status", "enum", ""),
        ("notes", "text?", ""),
        ("processed_by", "bigint?", "FK"),
        ("created_at", "timestamp?", ""),
        ("updated_at", "timestamp?", ""),
    ]

    users_box = draw_entity(draw, 80, 185, 520, "users", users, "#D45520")
    customers_box = draw_entity(draw, 80, 955, 520, "customers", customers, "#0F8B8D")
    bookings_box = draw_entity(draw, 910, 310, 650, "bookings", bookings, "#FF6626")
    services_box = draw_entity(draw, 2050, 185, 560, "services", services, "#4776B5")
    payments_box = draw_entity(draw, 1950, 895, 660, "payments", payments, "#7652A8")

    draw_relation(
        draw,
        (users_box[2], 390),
        (bookings_box[0], 465),
        "creates",
        "0..1",
        "0..N",
        bend=[(740, 390), (740, 465)],
    )
    draw_relation(
        draw,
        (customers_box[2], 1110),
        (bookings_box[0], 1060),
        "places",
        "0..1",
        "0..N",
        bend=[(745, 1110), (745, 1060)],
    )
    draw_relation(
        draw,
        (services_box[0], 520),
        (bookings_box[2], 570),
        "selected for",
        "1",
        "0..N",
        bend=[(1800, 520), (1800, 570)],
    )
    draw_relation(
        draw,
        (bookings_box[2], 1110),
        (payments_box[0], 1110),
        "has payment",
        "1",
        "0..1",
    )
    draw_relation(
        draw,
        (users_box[0] + 260, users_box[3]),
        (customers_box[0] + 260, customers_box[1]),
        "owns profile",
        "0..1",
        "0..1",
    )
    draw_relation(
        draw,
        (users_box[2], 650),
        (payments_box[0], 1380),
        "processed by",
        "0..1",
        "0..N",
        bend=[(730, 650), (730, 1560), (1780, 1560), (1780, 1380)],
    )

    legend_font = get_font(20)
    draw.rounded_rectangle((1740, 1580, 2610, 1660), radius=15, fill="#FFF9F1", outline="#E8DCCB", width=2)
    draw.text(
        (1765, 1606),
        "? = nullable  |  PK = primary key  |  FK = foreign key  |  UK = unique key",
        fill="#6B625A",
        font=legend_font,
    )
    image.save(DIAGRAM_PNG, quality=95)


def add_diagram_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(section, landscape=True)
    add_running_furniture(section)

    heading = doc.add_heading("3. Diagram ERD Inti", level=1)
    heading.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph(
        "Diagram berikut menampilkan entitas bisnis utama. Kardinalitas mengikuti relasi Eloquent dan alur aplikasi, "
        "dengan pengecualian constraint fisik yang dijelaskan pada bagian sebelumnya."
    )
    p.paragraph_format.space_after = Pt(5)

    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_after = Pt(0)
    picture = image_p.add_run().add_picture(str(DIAGRAM_PNG), width=Inches(9.75))
    picture._inline.docPr.set(
        "descr",
        "Diagram ERD VAULTLAUNDRY yang menampilkan relasi users, customers, services, bookings, dan payments.",
    )
    picture._inline.docPr.set("title", "Core Entity Relationship Diagram VAULTLAUNDRY")


def add_business_flow_and_notes(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(section, landscape=False)
    add_running_furniture(section)

    doc.add_heading("4. Alur Data Utama", level=1)
    steps = [
        ("Profil pelanggan", "User dapat memiliki profil customer. Kasir juga dapat mencatat customer walk-in tanpa akun user."),
        ("Pembuatan booking", "Booking mengacu ke satu service, lalu menyimpan tanggal, berat, pickup type, total harga, dan status proses."),
        ("Monitoring laundry", "Status booking bergerak dari booking_masuk sampai diambil atau dibatalkan."),
        ("Pembayaran", "Payment mengambil total_bill dari booking dan menyimpan metode, nominal, status pembayaran, serta processed_by."),
        ("Invoice dan laporan", "Invoice PDF dan laporan dibentuk dari relasi payment -> booking -> customer/service, tanpa tabel invoices atau reports."),
    ]
    for number, (title, description) in enumerate(steps, 1):
        table = doc.add_table(rows=1, cols=2)
        left, right = table.rows[0].cells
        set_cell_shading(left, ORANGE)
        set_cell_shading(right, CREAM)
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(number))
        set_repeatable_font(run, size=14, bold=True, color=WHITE)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(title)
        set_repeatable_font(run, size=10.5, bold=True, color=INK)
        p = right.add_paragraph(description)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_repeatable_font(run, size=9.5, color=MUTED)
        set_fixed_table_layout(table, [0.55, 6.1])
        set_table_borders(table)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    doc.add_heading("5. Aturan Integritas Data", level=1)
    rules = [
        "users.email, bookings.booking_code, dan payments.payment_code memiliki unique constraint.",
        "Penghapusan user mengubah customers.user_id, bookings.user_id, dan payments.processed_by menjadi null.",
        "Penghapusan customer mengubah bookings.customer_id menjadi null.",
        "Penghapusan service menghapus booking terkait melalui cascadeOnDelete.",
        "Penghapusan booking menghapus payment terkait melalui cascadeOnDelete.",
        "Status booking dibatasi ke delapan nilai enum; status payment dibatasi ke unpaid, partial, atau paid.",
    ]
    for rule in rules:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rule)

    doc.add_heading("6. Entitas Turunan yang Tidak Memiliki Tabel", level=1)
    table = doc.add_table(rows=1, cols=3)
    for i, header in enumerate(("Fitur", "Sumber Data", "Keterangan")):
        table.cell(0, i).text = header
        set_cell_shading(table.cell(0, i), ORANGE)
    derived = [
        ("Invoice PDF", "payments, bookings, customers, services, users", "Dibentuk saat diminta; tidak disimpan sebagai tabel invoices."),
        ("Laporan transaksi", "payments + relasi booking", "Hasil query dan filter; tidak disimpan sebagai tabel reports."),
        ("Laporan pendapatan", "payments", "Agregasi total paid dan piutang unpaid/partial."),
        ("Dashboard statistik", "customers, services, bookings, payments", "Agregasi realtime; bukan entitas database baru."),
    ]
    for values in derived:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_fixed_table_layout(table, [1.25, 2.65, 2.75])
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    style_table_text(table)


def add_mermaid_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Lampiran A. Kode Mermaid ERD", level=1)
    p = doc.add_paragraph(
        "Kode berikut dapat ditempel ke Mermaid Live Editor, Markdown yang mendukung Mermaid, "
        "atau dokumentasi GitHub. File sumber juga tersedia di database/erd-vaultlaundry.mmd."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    code_table = doc.add_table(rows=1, cols=1)
    cell = code_table.cell(0, 0)
    set_cell_shading(cell, "F5F2EE")
    code = MERMAID_FILE.read_text(encoding="utf-8")
    paragraph = cell.paragraphs[0]
    paragraph.style = doc.styles["ERD Code"]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(code)
    set_repeatable_font(run, name=MONO_FONT, size=7.3, color=INK)
    set_fixed_table_layout(code_table, [6.65])
    set_table_borders(code_table, color="D6C9BA")

    doc.add_heading("Lampiran B. Rekomendasi Penggunaan", level=1)
    recommendations = [
        "Gunakan Mermaid sebagai dokumentasi utama karena file teks mudah disimpan di Git dan mudah diperbarui bersama migration.",
        "Gunakan diagram di Word untuk laporan tugas, presentasi, atau dokumen akademik.",
        "Gunakan dbdiagram.io hanya jika membutuhkan pengaturan posisi tabel secara manual atau ekspor visual tambahan.",
        "Jika aplikasi benar-benar membutuhkan relasi 1:1 yang dijamin database, buat migration baru untuk unique index customers.user_id dan payments.booking_id.",
    ]
    for item in recommendations:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build_document() -> None:
    generate_diagram()
    doc = Document()
    apply_document_styles(doc)
    first_section = doc.sections[0]
    setup_section(first_section, landscape=False)
    first_section.different_first_page_header_footer = True
    add_running_furniture(first_section)

    core_properties = doc.core_properties
    core_properties.title = "Rancangan ERD VAULTLAUNDRY"
    core_properties.subject = "Entity Relationship Diagram aplikasi laundry berbasis Laravel"
    core_properties.author = "VAULTLAUNDRY Project"
    core_properties.keywords = "ERD, Laravel, Laundry, Database, Mermaid"
    core_properties.comments = "Dibuat berdasarkan migration dan model aktif."

    add_title_block(doc)
    add_overview(doc)
    add_entity_sections(doc)
    add_relationships(doc)
    add_diagram_section(doc)
    add_business_flow_and_notes(doc)
    add_mermaid_appendix(doc)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    print(DIAGRAM_PNG)


if __name__ == "__main__":
    build_document()
